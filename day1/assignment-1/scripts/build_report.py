#!/usr/bin/env python3
"""Build Assignment 1 REPORT.docx using Assignment 2 as the style template."""

from __future__ import annotations

import json
import math
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HERE = Path(__file__).resolve().parent
ASSIGNMENT = HERE.parent
ROOT = ASSIGNMENT.parent
IMGS = ASSIGNMENT / "imgs"
EVIDENCE = ASSIGNMENT / "evidence"
TEMPLATE = ROOT / "assignment-2" / "REPORT.docx"
OUTPUT = ASSIGNMENT / "REPORT.docx"

sys.path.insert(0, str(HERE))
from parse_bitcoin import Reader, decode_script, parse_block, parse_tx, read_hex_file  # noqa: E402


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=50, start=70, bottom=50, end=70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def format_cell(cell, font_size: float | None = None, bold: bool = False, monospace: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.bold = bold or run.bold
            if font_size:
                run.font.size = Pt(font_size)
            if monospace:
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_table(document: Document, rows: list[list[str]], widths: list[float] | None = None, font_size: float = 9.0):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table"
    table.autofit = False
    for r_index, values in enumerate(rows):
        for c_index, value in enumerate(values):
            cell = table.cell(r_index, c_index)
            cell.text = str(value)
            if widths:
                set_cell_width(cell, widths[c_index])
            if r_index == 0:
                set_cell_shading(cell, "D9EAF2")
            format_cell(cell, font_size=font_size, bold=r_index == 0)
    set_repeat_table_header(table.rows[0])
    return table


def add_paragraph(document: Document, text: str = "", style: str = "Body Text", bold_prefix: str | None = None):
    paragraph = document.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        first.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    return paragraph


def add_centered_formula(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Body Text")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    run.font.size = Pt(11)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)


def add_image(document: Document, filename: str, caption: str, width: float = 5.9) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(IMGS / filename), width=Inches(width))
    add_caption(document, caption)


def add_two_images(document: Document, items: list[tuple[str, str]], widths: tuple[float, float] = (2.85, 2.85)) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    remove_table_borders(table)
    for index, (filename, caption) in enumerate(items):
        cell = table.cell(0, index)
        set_cell_width(cell, 3.0)
        set_cell_margins(cell, top=30, start=40, bottom=30, end=40)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(IMGS / filename), width=Inches(widths[index]))
        cp = cell.add_paragraph(style="Caption")
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(caption)


def wrap_hex(value: str, width: int = 64) -> str:
    return "\n".join(value[i : i + width] for i in range(0, len(value), width))


def main() -> None:
    metadata = json.loads((EVIDENCE / "metadata.json").read_text(encoding="utf-8"))
    tx_data = read_hex_file(ASSIGNMENT / "tx.txt")
    tx_reader = Reader(tx_data)
    tx = parse_tx(tx_reader, "tx")
    block = parse_block(read_hex_file(EVIDENCE / "block-144878.hex"))
    fee = metadata["fee_sats"]
    rate = metadata["fee_rate_exact_sat_vb"]
    local_time = datetime.fromtimestamp(block["timestamp"], timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

    document = Document(TEMPLATE)
    clear_body(document)
    document.core_properties.title = "Bitcoin Testnet4 交易构造、逐字节解析与区块验证"
    document.core_properties.subject = "网络安全创新创业课程作业 1"
    document.core_properties.author = "魏子彦"
    document.core_properties.keywords = "Bitcoin; Testnet4; SegWit; P2WPKH; BIP143; transaction; block; Merkle"

    title = document.add_paragraph(style="Heading 1")
    title.add_run("Bitcoin Testnet4 交易构造、逐字节解析与区块验证")
    course = document.add_paragraph(style="First Paragraph")
    course.add_run("网络安全创新创业课程作业 1").bold = True
    document.add_paragraph(style="Normal")

    add_table(
        document,
        [
            ["项目", "内容"],
            ["学生姓名", "魏子彦"],
            ["学号", "202300460073"],
            ["专业/班级", "网安2班"],
            ["实验网络", "Bitcoin Testnet4"],
            ["交易 txid", metadata["txid"]],
            ["确认区块", f"高度 {metadata['block_height']}，{metadata['block_hash']}"],
            ["完成日期", "2026 年 7 月 20 日"],
        ],
        widths=[1.25, 4.75],
        font_size=9.5,
    )

    document.add_heading("摘要", level=2)
    add_paragraph(
        document,
        "本报告完成课件要求的 Bitcoin 实践项目：使用 Sparrow Wallet 在 Bitcoin Testnet4 创建原生 SegWit（P2WPKH）钱包，从测试币水龙头获得 UTXO，构造、签名并广播一笔交易；随后以最终广播的原始交易 tx.txt 为唯一权威数据源，对 222 个字节逐字段解释，并用脚本验证 txid、wtxid、交易重量、手续费、BIP143 签名哈希、HASH160 公钥绑定和 secp256k1 ECDSA 签名。该交易已进入高度 144878 的 Testnet4 区块，脚本进一步解析完整 4045 字节区块，覆盖 80 字节区块头和全部 14 笔交易，并重算 PoW 目标、交易 Merkle 根与 SegWit witness commitment。",
        style="First Paragraph",
    )
    add_paragraph(
        document,
        "最终交易 txid 为 1dd070a132e4c5dd1ab2fde424b82903a53e6afa22b921d9448590b71a5bcda7，花费水龙头 UTXO 394,027 sats，向目标地址支付 114,514 sats，找零 277,258 sats，实际手续费 2,255 sats。交易重量为 561 WU，精确虚拟大小为 140.25 vB，因此最终费率约为 16.08 sat/vB，与实际使用的 16× 设置一致。所有自动核验项均通过。",
    )
    add_paragraph(document, "关键词： Bitcoin；Testnet4；UTXO；SegWit；P2WPKH；BIP143；ECDSA；Merkle Tree；Proof of Work", bold_prefix="关键词：")

    document.add_heading("1. 作业目标与完成情况", level=2)
    add_paragraph(
        document,
        "课件要求在 Bitcoin 测试网上发送一笔交易，将原始交易数据解析到每一个字节，并尝试用脚本解析完整区块、计算各字段。为使结果可核验，本报告把任务拆为交易流程、序列化解析、脚本执行、密码学验签和区块验证五部分。",
        style="First Paragraph",
    )
    add_table(
        document,
        [
            ["要求", "完成方式", "核验结果"],
            ["测试网发送交易", "Sparrow Wallet 切换 Testnet4，接收测试币后构造、签名、广播", "已在高度 144878 确认"],
            ["逐字节解析交易", "脚本记录每个字段的绝对偏移、长度、原始十六进制和解释", "222/222 字节恰好覆盖一次"],
            ["解释锁定/解锁脚本", "解析前序 P2WPKH witness program、空 scriptSig、签名与压缩公钥 witness", "公钥哈希匹配，ECDSA 验签通过"],
            ["解析完整区块", "解析区块头、CompactSize 交易数及全部 14 笔交易", "4045/4045 字节恰好覆盖一次"],
            ["重算关键密码学结果", "重算 PoW、txid/wtxid、Merkle root、witness commitment", "全部与链上数据一致"],
        ],
        widths=[1.2, 3.55, 1.25],
        font_size=8.7,
    )
    add_paragraph(
        document,
        "安全说明：过程截图中的助记词属于钱包秘密，即使钱包仅用于测试网，也不应出现在可提交或公开传播的报告中。因此图 4 使用脱敏副本；原截图仍保留在本地 imgs 目录中，不在报告正文展示其内容。",
    )

    document.add_heading("2. 实验环境与交易流程", level=2)
    document.add_heading("2.1 安装钱包并切换 Testnet4", level=3)
    add_paragraph(
        document,
        "实验使用 Sparrow Wallet 2.5.2。启动后通过 Tools → Restart In → Testnet4 切换到第四代 Bitcoin 测试网。Testnet4 与主网使用相同的交易、脚本和工作量证明数据结构，但测试币没有真实经济价值，适合完成交易构造实验。",
        style="First Paragraph",
    )
    add_two_images(document, [("1.sparrow.png", "图 1  下载 Sparrow Wallet"), ("2.testnet4.png", "图 2  切换至 Testnet4")])

    document.add_heading("2.2 创建原生 SegWit 钱包", level=3)
    add_paragraph(
        document,
        "创建名为 Para wallet 的软件钱包，策略为 Single Signature HD，脚本类型选择 Native SegWit（P2WPKH），密钥派生遵循 BIP39/BIP84。BIP84 的典型 Testnet 派生路径为 m/84'/1'/account'/change/address_index。钱包生成的 12 个助记词用于恢复私钥，必须离线备份并避免截屏外泄。",
        style="First Paragraph",
    )
    add_two_images(document, [("3.wallet.png", "图 3  创建钱包"), ("4.mnemonic-redacted.png", "图 4  助记词页面（已脱敏）")])
    add_image(document, "5,apply.png", "图 5  应用 Native SegWit（P2WPKH）钱包策略", width=5.8)

    document.add_heading("2.3 从水龙头获得测试 UTXO", level=3)
    add_paragraph(
        document,
        "钱包生成接收地址 tb1qmf694grlk9r6cyt296dt5wr50803cmxy6vn2p2。水龙头向该地址支付 394,027 sats；该资金在后续交易中作为唯一输入，引用前序交易 97afb394…337a6a 的 vout=1。",
        style="First Paragraph",
    )
    add_two_images(document, [("6.address.png", "图 6  获取接收地址"), ("7.faucet.png", "图 7  Testnet4 水龙头转账")])
    add_two_images(document, [("8.receive.png", "图 8  钱包收到 394,027 sats"), ("9.confirmed.png", "图 9  首次确认通知")], widths=(2.9, 2.6))

    document.add_heading("2.4 构造、签名与广播交易", level=3)
    add_paragraph(
        document,
        "交易向 tb1qerzrlxcfu24davlur5sqmgzzgsal6wusda40er 支付 114,514 sats，并把 277,258 sats 找零到钱包新地址 tb1qycpmmqv7evxsexkcsaz6kl730x84zamzcdgpaq。最终交易使用实际 16× 费率设置；过程截图中出现的旧费率报价只表示较早的界面状态，最终数值以 tx.txt 和链上交易为准。",
        style="First Paragraph",
    )
    add_image(document, "10.create_tx.png", "图 10  创建交易并设置收款地址与金额", width=5.8)
    add_two_images(document, [("11.finalize_tx.png", "图 11  完成待签名交易"), ("12.sign.png", "图 12  由 BIP39 软件钱包签名")])
    add_two_images(document, [("13.broadcast.png", "图 13  广播已签名交易"), ("14.tx.png", "图 14  最终原始交易和 txid")])

    document.add_heading("3. 最终交易概览", level=2)
    add_table(
        document,
        [
            ["项目", "数值"],
            ["txid", tx["txid"]],
            ["wtxid", tx["wtxid"]],
            ["版本 / nLockTime", f"2 / {tx['locktime']}（区块高度语义）"],
            ["输入 / 输出", f"{len(tx['inputs'])} / {len(tx['outputs'])}"],
            ["总大小 / 基础大小 / witness", f"{tx['size']} / {tx['base_size']} / {tx['witness_size']} bytes"],
            ["重量 / vsize", f"{tx['weight']} WU / {tx['vsize']} vB（精确 {tx['exact_vsize']:.2f} vB）"],
            ["输入金额 / 输出合计", "394,027 / 391,772 sats"],
            ["手续费 / 实际费率", f"{fee:,} sats / {rate:.8f} sat/vB"],
            ["确认区块", f"高度 {metadata['block_height']}，交易索引 1"],
        ],
        widths=[1.55, 4.45],
        font_size=9.0,
    )
    add_centered_formula(document, "fee = 394,027 − 114,514 − 277,258 = 2,255 sats")
    add_centered_formula(document, "weight = base_size × 4 + witness_size = 113 × 4 + 109 = 561 WU")
    add_centered_formula(document, "fee rate = 2,255 / (561 / 4) ≈ 16.0784 sat/vB")
    add_table(
        document,
        [
            ["类型", "引用/序号", "金额", "脚本或地址"],
            ["输入", "97afb394…337a6a:1", "394,027 sats", "前序 P2WPKH：tb1qmf694…vn2p2"],
            ["输出 0（支付）", "vout=0", "114,514 sats", "0014c8c43f…bfd3b90 / tb1qerzrl…da40er"],
            ["输出 1（找零）", "vout=1", "277,258 sats", "00142603bd…517762 / tb1qycpmm…cdgpaq"],
        ],
        widths=[1.0, 1.45, 1.0, 2.55],
        font_size=8.6,
    )

    document.add_heading("4. 原始交易逐字节解析", level=2)
    document.add_heading("4.1 序列化顺序与大小端", level=3)
    add_paragraph(
        document,
        "Bitcoin 交易不是 JSON，而是紧凑的二进制串。整数通常按 little-endian 写入；交易哈希在原始数据中也按内部字节序出现，显示给用户时再反转。输入/输出数量和脚本长度使用 CompactSize。SegWit 交易在 version 后加入 marker=00、flag=01，并把 witness 放在全部输出之后、nLockTime 之前。",
        style="First Paragraph",
    )
    add_centered_formula(document, "SegWit tx = version || 00 || 01 || vin || vout || witness || nLockTime")
    add_centered_formula(document, "txid = reverse(SHA256d(stripped serialization))")
    add_centered_formula(document, "wtxid = reverse(SHA256d(full SegWit serialization))")
    add_paragraph(
        document,
        "本交易的 stripped serialization 不含 marker、flag 和 witness，共 113 字节；完整序列化为 222 字节。因此 txid 与 wtxid 不同。逐字节脚本检查确认 offset 0 到 221 连续覆盖，没有遗漏或重叠。",
    )

    document.add_heading("4.2 位级语义", level=3)
    add_paragraph(
        document,
        "交易序列化的最小寻址单位是字节；所谓“到每一 bit”主要体现在标志位、操作码和前缀。下表把本交易中有明确位语义的字段进一步展开。",
        style="First Paragraph",
    )
    add_table(
        document,
        [
            ["字段", "原始值", "二进制/位语义"],
            ["version", "02 00 00 00", "按 little-endian 解释为 32 位整数 2"],
            ["marker", "00", "00000000：提示采用扩展序列化"],
            ["flag", "01", "00000001：bit 0 表示存在 witness"],
            ["P2WPKH 长度", "16", "00010110₂ = 22 字节"],
            ["sequence", "fd ff ff ff", "0xfffffffd；bit31=1 禁用 BIP68 相对锁定，同时满足 BIP125 opt-in RBF"],
            ["witness 项数", "02", "00000010₂ = 2 项"],
            ["签名长度", "47", "01000111₂ = 71 字节"],
            ["公钥长度", "21", "00100001₂ = 33 字节"],
            ["公钥前缀", "02", "压缩公钥，表示 y 坐标为偶数"],
            ["sighash", "01", "SIGHASH_ALL；未设置 ANYONECANPAY"],
            ["nLockTime", "ec 35 02 00", "little-endian = 144876；小于 5×10⁸，按区块高度解释"],
        ],
        widths=[1.25, 1.25, 3.5],
        font_size=8.5,
    )

    document.add_heading("4.3 完整字段表", level=3)
    add_paragraph(
        document,
        "下表列出 tx.txt 中全部字段。空 scriptSig 的长度为 0，因此它占用 0 个数据字节；长度字段本身仍占 1 字节。长签名和公钥保持完整十六进制，不做省略。",
        style="First Paragraph",
    )
    field_rows = [["偏移", "字节数", "字段", "原始十六进制", "解释"]]
    for field in tx["fields"]:
        raw = field.raw.hex() if field.length else "<empty>"
        field_rows.append([f"{field.offset} / 0x{field.offset:04x}", str(field.length), field.path.replace("tx.", ""), raw, field.decoded])
    field_table = add_table(document, field_rows, widths=[0.72, 0.45, 1.2, 2.2, 1.43], font_size=6.8)
    for row_index, row in enumerate(field_table.rows):
        format_cell(row.cells[3], font_size=6.1, monospace=True, bold=row_index == 0)

    document.add_heading("5. P2WPKH 锁定与解锁脚本", level=2)
    document.add_heading("5.1 前序 UTXO 的锁定条件", level=3)
    add_paragraph(
        document,
        "本交易输入本身只保存“引用哪个 UTXO”，真正的锁定脚本来自被花费的前序输出。链上前序输出脚本为 0014da745aa07fb147ac116a2e9aba387479df1c6cc4：00 是 witness version 0，14 是 20 字节数据长度，后续 20 字节是公钥哈希。对应地址为 tb1qmf694grlk9r6cyt296dt5wr50803cmxy6vn2p2。",
        style="First Paragraph",
    )
    add_centered_formula(document, "scriptPubKey = OP_0 || PUSH20 || HASH160(expected public key)")

    document.add_heading("5.2 witness 解锁数据", level=3)
    add_paragraph(
        document,
        "由于是原生 P2WPKH，scriptSig 为空，解锁数据位于 witness。witness 有两项：第一项是 70 字节 DER 签名加 1 字节 SIGHASH_ALL，第二项是 33 字节压缩公钥。",
        style="First Paragraph",
    )
    add_table(
        document,
        [
            ["witness 项", "长度", "内容"],
            ["0", "71", tx["witnesses"][0][0].hex()],
            ["1", "33", tx["witnesses"][0][1].hex()],
        ],
        widths=[0.75, 0.65, 4.6],
        font_size=7.5,
    )
    add_paragraph(
        document,
        "验证先计算 HASH160(pubkey)。结果 da745aa07fb147ac116a2e9aba387479df1c6cc4 与前序 witness program 完全一致，证明 witness 公钥满足地址绑定；随后按 BIP143 计算签名消息并执行 CHECKSIG。",
    )

    document.add_heading("5.3 DER 签名与 BIP143 消息", level=3)
    add_paragraph(
        document,
        "签名结构为 30 44 02 20 [r] 02 20 [s] 01。30 表示 DER SEQUENCE，44 表示后续 DER 内容 68 字节；两个 02 20 分别表示 32 字节正整数 r 和 s；最后的 01 不属于 DER，而是 Bitcoin 的 SIGHASH_ALL 类型。",
        style="First Paragraph",
    )
    add_table(
        document,
        [
            ["项目", "数值"],
            ["r", "37462debb2cccf9d64623982195ea538e6ed24fe79b62aad41c328d49f1cb549"],
            ["s", "32954dd35fcaeb8fa015d442d0db092a2c6706e4d8433ced0edd62e46a0da691"],
            ["low-S", "是，s ≤ n/2，满足 Bitcoin 标准化要求"],
            ["scriptCode", "76a914da745aa07fb147ac116a2e9aba387479df1c6cc488ac"],
            ["BIP143 digest", "69cbe0b2ddf0d3b96bf4ecb9f9a9c5291ac6ddf55848ef91b5c6d18a3808cc83"],
            ["ECDSA 结果", "通过：Rₓ mod n = r"],
        ],
        widths=[1.15, 4.85],
        font_size=8.4,
    )
    add_centered_formula(document, "z = SHA256d(nVersion || hashPrevouts || hashSequence || outpoint || scriptCode || amount || nSequence || hashOutputs || nLockTime || hashType)")
    add_centered_formula(document, "R = s⁻¹(zG + rP),  verify if Rₓ mod n = r")
    add_paragraph(
        document,
        "脚本使用真实 secp256k1 曲线参数做纯 Python ECDSA 验证，不依赖钱包或区块浏览器的“已验证”结论。该校验通过，说明签名确由前序 UTXO 对应私钥产生，并且绑定了本交易的全部输入、sequence 和输出。",
    )

    document.add_heading("6. 完整区块解析与验证", level=2)
    document.add_heading("6.1 区块头和 PoW", level=3)
    add_paragraph(
        document,
        f"交易被打包进 Testnet4 高度 {metadata['block_height']} 的区块。完整区块原始数据为 {block['size']} 字节，其中区块头固定 80 字节，随后 1 字节 CompactSize 值 0e 表示 14 笔交易。区块时间为 {local_time}。",
        style="First Paragraph",
    )
    add_table(
        document,
        [
            ["偏移", "长度", "区块头字段", "原始值", "解释"],
            ["0", "4", "version", "0060bb31", "0x31bb6000"],
            ["4", "32", "previous block", "6d85af4f…00000000", block["previous_block_hash"]],
            ["36", "32", "Merkle root", "02d1ed18…b645358", block["stored_merkle_root"]],
            ["68", "4", "time", "55f05d6a", local_time],
            ["72", "4", "bits", "ffff001d", "0x1d00ffff"],
            ["76", "4", "nonce", "94023841", str(block["nonce"])],
            ["80", "1", "tx count", "0e", "14"],
        ],
        widths=[0.55, 0.45, 1.1, 1.35, 2.55],
        font_size=7.8,
    )
    add_centered_formula(document, "block_hash = reverse(SHA256d(80-byte header))")
    add_paragraph(document, f"重算区块哈希得到 {block['block_hash']}，与链上区块哈希一致。bits=0x1d00ffff 解码出的目标为 00000000ffff0000…0000；把区块哈希视为大端整数后满足 hash ≤ target，因此 PoW 校验通过。")

    document.add_heading("6.2 交易 Merkle 根", level=3)
    add_paragraph(
        document,
        "普通交易 Merkle 树以每笔交易的内部 txid 字节为叶子，相邻两项连接后做 SHA256d；当某层节点数为奇数时复制最后一项。14 个叶子最终得到的根为 5853640bbd50baf9df86ae513c514b856ba78e39933d2ea5fda3df2518edd102，与区块头完全一致。作业交易位于区块交易数组索引 1（coinbase 之后的第一笔普通交易）。",
        style="First Paragraph",
    )
    tx_rows = [["索引", "txid", "size / weight / vsize", "vin / vout"]]
    for index, item in enumerate(block["transactions"]):
        marker = "（本作业）" if item["txid"] == tx["txid"] else ""
        tx_rows.append([
            f"{index}{marker}",
            item["txid"],
            f"{item['size']} B / {item['weight']} WU / {item['vsize']} vB",
            f"{len(item['inputs'])} / {len(item['outputs'])}",
        ])
    tx_index_table = add_table(document, tx_rows, widths=[0.7, 3.15, 1.45, 0.7], font_size=6.9)
    for row_index, row in enumerate(tx_index_table.rows):
        format_cell(row.cells[1], font_size=6.2, monospace=True, bold=row_index == 0)

    document.add_heading("6.3 SegWit witness commitment", level=3)
    add_paragraph(
        document,
        "SegWit 区块还必须承诺 witness 数据。计算时把 coinbase 的 wtxid 设为 32 字节全零，其余交易使用 wtxid 构造 witness Merkle 树；再与 coinbase witness 中的 32 字节 reserved value 连接并 SHA256d。",
        style="First Paragraph",
    )
    add_centered_formula(document, "commitment = SHA256d(witness_merkle_root || witness_reserved_value)")
    add_table(
        document,
        [
            ["项目", "计算值"],
            ["witness Merkle root", block["witness_merkle_root"]],
            ["reserved value", block["witness_reserved_value"]],
            ["计算 commitment", block["calculated_witness_commitment"]],
            ["coinbase OP_RETURN 中的 commitment", block["stored_witness_commitment"]],
            ["结果", "完全一致，witness commitment 校验通过"],
        ],
        widths=[1.65, 4.35],
        font_size=8.2,
    )

    document.add_heading("6.4 完整区块字节覆盖", level=3)
    add_paragraph(
        document,
        f"脚本从 offset 0 开始顺序读取区块头、交易数和 14 笔交易的 version、vin、vout、witness、locktime，最终 reader offset 恰为 {block['size']}。基础大小为 {block['stripped_size']} 字节，总重量 {block['weight']} WU，vsize {block['vsize']} vB。evidence/block_parse.txt 保留完整 43 KB 字段清单，逐行给出 offset、长度、字段路径、完整十六进制和解释，从而证明 4045/4045 字节恰好覆盖一次。",
        style="First Paragraph",
    )
    add_paragraph(document, "Coinbase 的 BIP34 高度解码为 144878；其首个输出为 5,000,011,515 sats，即 50 BTC 区块补贴加本区块总手续费 11,515 sats。第二个 0 sat OP_RETURN 输出携带 witness commitment。")

    document.add_heading("7. 自动化脚本与复现结果", level=2)
    add_paragraph(
        document,
        "scripts/parse_bitcoin.py 仅使用 Python 标准库，包含 CompactSize、SegWit 交易、常见 scriptPubKey、DER、secp256k1 ECDSA、Merkle 树和 compact target 实现。它读取 tx.txt、metadata.json 和 block-144878.hex，生成两份证据报告。",
        style="First Paragraph",
    )
    code = document.add_paragraph(style="Source Code")
    code.add_run("cd assignment-1\npython .\\scripts\\parse_bitcoin.py")
    add_table(
        document,
        [
            ["自动核验项", "结果"],
            ["txid / wtxid", "PASS / PASS"],
            ["BIP143 ECDSA signature", "PASS"],
            ["block hash / BIP34 height", "PASS / PASS"],
            ["proof of work", "PASS"],
            ["transaction Merkle root", "PASS"],
            ["SegWit witness commitment", "PASS"],
            ["交易位于区块中", "PASS"],
            ["222 字节交易覆盖", "PASS"],
            ["4045 字节区块覆盖", "PASS"],
        ],
        widths=[3.4, 2.6],
        font_size=9.0,
    )
    add_paragraph(document, "evidence/tx_parse.txt 是交易解析与签名验证的可读输出；evidence/block_parse.txt 是完整区块字段表；evidence/block-144878.hex 保存可离线复现的区块原始数据。")

    document.add_heading("8. 结论", level=2)
    add_paragraph(
        document,
        "本实验从钱包界面一直追踪到共识数据结构。Sparrow Wallet 展示的“收款、找零、手续费”最终都落为确定的 little-endian 整数和脚本字节；地址不是链上字段，而是 witness program 的人类可读 Bech32 编码；P2WPKH 的所有权证明不在 scriptSig，而在 witness 的签名与公钥中；txid 为兼容旧系统排除 witness，wtxid 和 coinbase witness commitment 则把 witness 纳入区块承诺。",
        style="First Paragraph",
    )
    add_paragraph(
        document,
        "最终交易的 222 字节、所在区块的 4045 字节均得到无遗漏解析。交易金额守恒、BIP143 摘要、ECDSA 签名、PoW、交易 Merkle 根和 witness commitment 全部通过独立重算。由此可见，Bitcoin 的安全不是单一签名算法的结果，而是 UTXO 引用、脚本条件、序列化规则、哈希承诺和工作量证明共同构成的可验证链条。",
    )

    document.add_heading("参考资料", level=2)
    references = [
        "Satoshi Nakamoto, Bitcoin: A Peer-to-Peer Electronic Cash System, 2008.",
        "Bitcoin Developer Reference, Transactions and Block Chain data structures.",
        "BIP 34: Block v2, Height in Coinbase.",
        "BIP 125: Opt-in Full Replace-by-Fee Signaling.",
        "BIP 141: Segregated Witness (Consensus layer).",
        "BIP 143: Transaction Signature Verification for Version 0 Witness Program.",
        "BIP 173: Base32 address format for native v0-16 witness outputs.",
        "BIP 84: Derivation scheme for P2WPKH based accounts.",
        "Sparrow Wallet Documentation, https://sparrowwallet.com/docs/",
        "mempool.space Testnet4 API, transaction and block raw data.",
    ]
    for reference in references:
        add_paragraph(document, reference, style="Compact")

    document.add_heading("附录 A：最终原始交易", level=2)
    raw = document.add_paragraph(style="Source Code")
    raw.add_run(wrap_hex(tx_data.hex(), 64))

    document.add_heading("附录 B：证据文件索引", level=2)
    add_table(
        document,
        [
            ["文件", "用途"],
            ["tx.txt", "最终广播交易的 222 字节原始十六进制"],
            ["scripts/parse_bitcoin.py", "逐字节解析和密码学验证脚本"],
            ["evidence/metadata.json", "链上前序输出、确认区块和最终费率元数据"],
            ["evidence/tx_parse.txt", "完整交易字段表、BIP143 摘要和 ECDSA 结果"],
            ["evidence/block-144878.hex", "包含该交易的完整 Testnet4 区块原始数据"],
            ["evidence/block_parse.txt", "完整区块逐字段解析与 PoW/Merkle/witness 验证"],
            ["imgs/", "实验过程截图；助记词副本在报告中已脱敏"],
        ],
        widths=[2.3, 3.7],
        font_size=8.8,
    )

    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")
    print(f"paragraphs={len(document.paragraphs)}, tables={len(document.tables)}")


if __name__ == "__main__":
    main()
